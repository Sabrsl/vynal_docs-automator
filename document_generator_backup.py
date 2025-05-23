import os
import re
import logging
from datetime import datetime
import io

# Pour les documents PDF
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm, inch
    from reportlab.platypus import Paragraph, Table, TableStyle, Spacer, Image
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from PIL import Image as PILImage
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("ReportLab n'est pas installÃ©. La gÃ©nÃ©ration de PDF sera limitÃ©e.")

# Pour les documents Word
try:
    import docx
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx n'est pas installÃ©. La gÃ©nÃ©ration de DOCX sera limitÃ©e.")

logger = logging.getLogger("VynalDocsAutomator.DocumentGenerator")

class DocumentGenerator:
    """
    GÃ©nÃ©rateur de documents pour l'application
    Version amÃ©liorÃ©e avec gestion d'erreur robuste
    """
    
    def __init__(self, config_manager=None):
        """
        Initialise le gÃ©nÃ©rateur de documents
        
        Args:
            config_manager: Gestionnaire de configuration
        """
        self.config = config_manager
        logger.info("DocumentGenerator initialisÃ© avec les capacitÃ©s suivantes:")
        logger.info(f"- PDF (ReportLab): {'Disponible' if REPORTLAB_AVAILABLE else 'Non disponible'}")
        logger.info(f"- DOCX (python-docx): {'Disponible' if DOCX_AVAILABLE else 'Non disponible'}")
    
    def strip_html(self, html_content):
        """
        EnlÃ¨ve les balises HTML du contenu tout en prÃ©servant le texte
        et extrait les images pour le traitement ultÃ©rieur
        
        Args:
            html_content: Contenu HTML
            
        Returns:
            tuple: (texte sans balises HTML, liste d'images extraites)
        """
        # Images extraites
        images = []
        
        # Si le contenu n'est pas une chaÃ®ne, le convertir
        if not isinstance(html_content, str):
            logger.warning(f"Le contenu n'est pas une chaÃ®ne, conversion en: {str(html_content)}")
            return str(html_content), images
        
        # Journaliser le contenu HTML pour le dÃ©bogage
        logger.info(f"Contenu HTML Ã  traiter (dÃ©but): {html_content[:100]}...")
        
        # Calculer la position relative des images dans le contenu
        content_length = len(html_content)
        
        # Extraire les balises d'image pour traitement ultÃ©rieur
        img_pattern = r'<img\s+[^>]*src\s*=\s*["\']([^"\']+)["\'][^>]*>'
        img_matches = list(re.finditer(img_pattern, html_content))
        logger.info(f"Nombre d'images trouvÃ©es dans le HTML: {len(img_matches)}")
        
        for img_match in img_matches:
            img_tag = img_match.group(0)  # RÃ©cupÃ©rer la balise img complÃ¨te
            img_src = img_match.group(1)
            # Journaliser pour le dÃ©bogage
            logger.info(f"Image trouvÃ©e avec src: {img_src[:50]}...")
            
            # Extraire les dimensions si disponibles - vÃ©rifier plusieurs formats possibles
            width_match = re.search(r'width\s*=\s*["\'](\d+)["\']', img_tag)
            if not width_match:
                width_match = re.search(r'width\s*=\s*(\d+)', img_tag)
                
            height_match = re.search(r'height\s*=\s*["\'](\d+)["\']', img_tag)
            if not height_match:
                height_match = re.search(r'height\s*=\s*(\d+)', img_tag)
            
            # Extraire l'alignement si disponible
            align_match = re.search(r'align\s*=\s*["\']([^"\']+)["\']', img_tag)
            if not align_match:
                align_match = re.search(r'style\s*=\s*["\'][^"\']*text-align\s*:\s*([^;"\'\s]+)[^"\']*["\']', img_tag)
            
            # Calculer la position relative dans le document (0 Ã  1)
            relative_position = img_match.start() / max(1, content_length)
            
            # Convertir les dimensions en entiers si trouvÃ©es
            width = None
            height = None
            align = None
            
            if width_match:
                try:
                    width = int(width_match.group(1))
                    logger.info(f"Largeur extraite: {width}")
                except ValueError:
                    logger.warning(f"Impossible de convertir la largeur en entier: {width_match.group(1)}")
            
            if height_match:
                try:
                    height = int(height_match.group(1))
                    logger.info(f"Hauteur extraite: {height}")
                except ValueError:
                    logger.warning(f"Impossible de convertir la hauteur en entier: {height_match.group(1)}")
            
            if align_match:
                align = align_match.group(1).lower()
                logger.info(f"Alignement de l'image: {align}")
            
            # Si c'est une image en base64, la sauvegarder dans un fichier temporaire
            if img_src.startswith('data:image/'):
                try:
                    import base64
                    import tempfile
                    
                    # Extraire le type et les donnÃ©es
                    header, data = img_src.split(',', 1)
                    mime_type = header.split(';')[0].split(':')[1]
                    img_ext = mime_type.split('/')[1]
                    
                    # CrÃ©er un fichier temporaire
                    temp_fd, temp_img_path = tempfile.mkstemp(suffix=f'.{img_ext}')
                    os.close(temp_fd)
                    
                    # DÃ©coder et sauvegarder l'image
                    img_data = base64.b64decode(data)
                    with open(temp_img_path, 'wb') as f:
                        f.write(img_data)
                    
                    # Utiliser ce chemin temporaire au lieu de l'URL base64
                    img_src = temp_img_path
                    logger.info(f"Image base64 convertie en fichier temporaire: {temp_img_path}")
                except Exception as e:
                    logger.warning(f"Erreur lors du traitement de l'image base64: {e}")
            
            # Ajouter l'image Ã  la liste avec sa position et dimensions
            image_info = {
                'src': img_src,
                'position': img_match.start(),
                'relative_position': relative_position
            }
            if width is not None:
                image_info['width'] = width
                logger.info(f"Largeur de l'image ajoutÃ©e: {width}")
            if height is not None:
                image_info['height'] = height
                logger.info(f"Hauteur de l'image ajoutÃ©e: {height}")
            if align is not None:
                image_info['align'] = align
            
            images.append(image_info)
        
        # Supprimer les balises d'image mais ajouter un marqueur pour leur emplacement
        html_content = re.sub(img_pattern, '[IMAGE]', html_content)
        
        # Compter les occurrences de [IMAGE] pour vÃ©rifier la cohÃ©rence
        image_markers = html_content.count('[IMAGE]')
        logger.info(f"Nombre de marqueurs [IMAGE] dans le texte: {image_markers}")
        if image_markers != len(images):
            logger.warning(f"IncohÃ©rence: {len(images)} images extraites mais {image_markers} marqueurs [IMAGE]")
        
        # Traiter certaines balises spÃ©ciales pour conserver la mise en forme
        # Remplacer les listes Ã  puces par des caractÃ¨res spÃ©ciaux
        html_content = re.sub(r'<li>(.*?)</li>', r'â€¢ \1\n', html_content)
        
        # Remplacer les sauts de ligne HTML par des sauts de ligne rÃ©els
        html_content = html_content.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
        
        # Remplacer les paragraphes par des lignes avec double saut de ligne
        html_content = re.sub(r'<p>(.*?)</p>', r'\1\n\n', html_content)
        
        # Remplacer les titres avec mise en forme
        html_content = re.sub(r'<h1>(.*?)</h1>', r'\1\n\n', html_content)
        html_content = re.sub(r'<h2>(.*?)</h2>', r'\1\n\n', html_content)
        html_content = re.sub(r'<h3>(.*?)</h3>', r'\1\n\n', html_content)
        
        # Remplacer les balises de style courantes
        html_content = re.sub(r'<strong>(.*?)</strong>', r'\1', html_content)
        html_content = re.sub(r'<b>(.*?)</b>', r'\1', html_content)
        html_content = re.sub(r'<em>(.*?)</em>', r'\1', html_content)
        html_content = re.sub(r'<i>(.*?)</i>', r'\1', html_content)
        html_content = re.sub(r'<u>(.*?)</u>', r'\1', html_content)
        
        # Supprimer les balises de listes
        html_content = html_content.replace('<ul>', '').replace('</ul>', '\n')
        html_content = html_content.replace('<ol>', '').replace('</ol>', '\n')
        
        # Supprimer toutes les autres balises HTML
        html_content = re.sub(r'<[^>]*>', '', html_content)
        
        # Supprimer les espaces multiples
        html_content = re.sub(r' +', ' ', html_content)
        
        # Supprimer les sauts de ligne multiples
        html_content = re.sub(r'\n{3,}', '\n\n', html_content)
        
        # GÃ©rer les entitÃ©s HTML courantes et avancÃ©es
        html_entities = {
            '&amp;': '&', 
            '&lt;': '<', 
            '&gt;': '>', 
            '&quot;': '"', 
            '&apos;': "'",
            '&nbsp;': ' ',
            '&copy;': 'Â©',
            '&reg;': 'Â®',
            '&trade;': 'â„¢',
            '&euro;': 'â‚¬',
            '&pound;': 'Â£',
            '&yen;': 'Â¥',
            '&cent;': 'Â¢',
            '&deg;': 'Â°',
            '&plusmn;': 'Â±',
            '&divide;': 'Ã·',
            '&times;': 'Ã—',
            '&sect;': 'Â§',
            '&para;': 'Â¶',
            '&micro;': 'Âµ',
            '&middot;': 'Â·',
            '&bull;': 'â€¢',
            '&hellip;': 'â€¦',
            '&ndash;': 'â€“',
            '&mdash;': 'â€”',
            '&lsquo;': ''',
            '&rsquo;': ''',
            '&ldquo;': '"',
            '&rdquo;': '"',
            '&laquo;': 'Â«',
            '&raquo;': 'Â»',
            '&frac14;': 'Â¼',
            '&frac12;': 'Â½',
            '&frac34;': 'Â¾',
            '&larr;': 'â†',
            '&uarr;': 'â†‘',
            '&rarr;': 'â†’',
            '&darr;': 'â†“',
            '&infin;': 'âˆž',
            '&ne;': 'â‰ ',
            '&asymp;': 'â‰ˆ',
            '&le;': 'â‰¤',
            '&ge;': 'â‰¥',
            '&sum;': 'âˆ‘',
            '&int;': 'âˆ«',
            '&alpha;': 'Î±',
            '&beta;': 'Î²',
            '&gamma;': 'Î³',
            '&delta;': 'Î´',
            '&epsilon;': 'Îµ',
            '&theta;': 'Î¸',
            '&lambda;': 'Î»',
            '&mu;': 'Î¼',
            '&pi;': 'Ï€',
            '&sigma;': 'Ïƒ',
            '&phi;': 'Ï†',
            '&omega;': 'Ï‰'
        }
        
        for entity, char in html_entities.items():
            html_content = html_content.replace(entity, char)
        
        # Traiter les entitÃ©s HTML numÃ©riques (dÃ©cimales et hexadÃ©cimales)
        def replace_numeric_entity(match):
            try:
                value = match.group(1)
                if value.startswith('x'):
                    # EntitÃ© hexadÃ©cimale
                    code = int(value[1:], 16)
                else:
                    # EntitÃ© dÃ©cimale
                    code = int(value)
                return chr(code)
            except:
                return match.group(0)  # Retourner l'entitÃ© originale en cas d'erreur
        
        # Remplacer les entitÃ©s numÃ©riques
        html_content = re.sub(r'&#(\d+);', lambda m: chr(int(m.group(1))), html_content)
        html_content = re.sub(r'&#x([0-9a-fA-F]+);', lambda m: chr(int(m.group(1), 16)), html_content)
        
        return html_content.strip(), images
    
    def replace_variables(self, content, variables):
        """
        Remplace les variables dans un contenu de maniÃ¨re sÃ©curisÃ©e
        Supporte Ã  la fois les formats {variable} et {{variable}}
        
        Args:
            content: Contenu avec variables
            variables: Dictionnaire des variables et leurs valeurs
            
        Returns:
            str: Contenu avec variables remplacÃ©es
        """
        import re
        
        # Si le contenu n'est pas une chaÃ®ne, le convertir
        if not isinstance(content, str):
            content = str(content)
        
        # PrÃ©paration des valeurs pour Ã©viter les valeurs None
        safe_variables = {}
        for key, value in variables.items():
            if value is None:
                safe_variables[key] = ""
            else:
                safe_variables[key] = str(value)
        
        # Remplacer d'abord les variables au format {{variable}}
        for var_name, var_value in safe_variables.items():
            # Format {{variable}}
            pattern = r"{{" + re.escape(var_name) + r"}}"
            content = re.sub(pattern, var_value, content)
        
        # Ensuite remplacer les variables au format {variable}
        for var_name, var_value in safe_variables.items():
            # Format {variable} (pour la rÃ©trocompatibilitÃ©)
            # Utiliser une expression rÃ©guliÃ¨re pour Ã©viter les faux positifs
            pattern = r"{" + re.escape(var_name) + r"}"
            content = re.sub(pattern, var_value, content)
        
        # Rechercher les variables non remplacÃ©es au format {{variable}}
        remaining_vars = re.findall(r'{{([^{}]*?)}}', content)
        if remaining_vars:
            logger.warning(f"Variables non remplacÃ©es (format {{variable}}): {remaining_vars}")
            # Remplacer les variables non trouvÃ©es par une chaÃ®ne vide
            for var in remaining_vars:
                content = content.replace(f"{{{{{var}}}}}", "")
        
        # Rechercher les variables non remplacÃ©es au format {variable}
        remaining_simple_vars = re.findall(r'{([^{}]*)}', content)
        if remaining_simple_vars:
            logger.warning(f"Variables non remplacÃ©es (format {variable}): {remaining_simple_vars}")
            # Remplacer les variables non trouvÃ©es par une chaÃ®ne vide
            for var in remaining_simple_vars:
                content = content.replace(f"{{{var}}}", "")
        
        return content
    
    def clean_filename(self, name):
        """
        Nettoie un nom pour qu'il soit utilisable dans un nom de fichier
        
        Args:
            name: Nom Ã  nettoyer
            
        Returns:
            str: Nom nettoyÃ©
        """
        # Supprimer les caractÃ¨res spÃ©ciaux et remplacer les espaces par des underscores
        name = re.sub(r'[\\/*?:"<>|]', '', name)
        name = name.replace(' ', '_')
        name = name.replace('/', '_')
        name = name.replace('\\', '_')
        name = name.strip('.')  # Supprime les points en dÃ©but et fin
        
        return name
    
    def generate_document(self, file_path, template, client, company_info, variables, format_type=None):
        """
        GÃ©nÃ¨re un document Ã  partir d'un modÃ¨le avec gestion d'erreur robuste
        
        Args:
            file_path: Chemin du fichier Ã  crÃ©er
            template: ModÃ¨le de document (dict)
            client: Informations du client (dict)
            company_info: Informations de l'entreprise (dict)
            variables: Variables spÃ©cifiques pour le document
            format_type: Format du document (pdf, docx ou txt)
            
        Returns:
            bool: True si le document a Ã©tÃ© gÃ©nÃ©rÃ© avec succÃ¨s, False sinon
        """
        try:
            # DÃ©terminer le format si non spÃ©cifiÃ©
            if format_type is None:
                format_type = "pdf"
                if self.config:
                    format_type = self.config.get("document.default_format", "pdf")
            
            format_type = format_type.lower()
            
            # VÃ©rifier la disponibilitÃ© des modules nÃ©cessaires
            if format_type == "pdf" and not REPORTLAB_AVAILABLE:
                logger.warning("ReportLab n'est pas installÃ©, utilisation du format TXT Ã  la place")
                format_type = "txt"
            elif format_type == "docx" and not DOCX_AVAILABLE:
                logger.warning("python-docx n'est pas installÃ©, utilisation du format TXT Ã  la place")
                format_type = "txt"
            
            # S'assurer que l'extension correspond au format
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext[1:] != format_type:
                file_path = os.path.splitext(file_path)[0] + f".{format_type}"
            
            # S'assurer que le rÃ©pertoire parent existe
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # PrÃ©parer les variables
            all_variables = {}
            
            # Variables standard du client
            all_variables.update({
                "client_name": client.get("name", ""),
                "client_company": client.get("company", ""),
                "client_email": client.get("email", ""),
                "client_phone": client.get("phone", ""),
                "client_address": client.get("address", "")
            })
            
            # Variables de l'entreprise
            all_variables.update({
                "company_name": company_info.get("name", ""),
                "company_address": company_info.get("address", ""),
                "company_email": company_info.get("email", ""),
                "company_phone": company_info.get("phone", ""),
                "company_website": company_info.get("website", "")
            })
            
            # Variables spÃ©cifiques
            all_variables.update(variables)
            
            # Date actuelle
            all_variables["date"] = datetime.now().strftime("%Y-%m-%d")
            
            # Obtenir le contenu du modÃ¨le
            content = template.get("content", "")
            
            # Journaliser la taille du contenu pour le dÃ©bogage
            content_sample = content[:100] + "..." if len(content) > 100 else content
            logger.info(f"Contenu du modÃ¨le (Ã©chantillon): {content_sample}")
            logger.info(f"Taille du contenu du modÃ¨le: {len(content)} caractÃ¨res")
            logger.info(f"Le contenu contient des balises img: {'<img' in content}")
            
            # Si le modÃ¨le a un chemin de fichier, l'utiliser
            if "file_path" in template and template["file_path"]:
                try:
                    template_file = template["file_path"]
                    if not os.path.isabs(template_file):
                        # Construire le chemin absolu si le chemin est relatif
                        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                        template_file = os.path.join(base_dir, "data", "templates", template_file)
                    
                    if os.path.exists(template_file):
                        with open(template_file, 'r', encoding='utf-8') as f:
                            content = f.read()
                except Exception as e:
                    logger.error(f"Erreur lors de la lecture du fichier modÃ¨le: {e}")
                    # Continuer avec le contenu stockÃ© dans la base de donnÃ©es
            
            # Remplacer les variables dans le contenu
            content = self.replace_variables(content, all_variables)
            
            # VÃ©rifier si le contenu aprÃ¨s remplacement des variables contient toujours des balises d'image
            logger.info(f"AprÃ¨s remplacement des variables, contient des balises img: {'<img' in content}")
            if '<img' in content:
                # Compter les balises d'image
                img_count = content.count('<img')
                logger.info(f"Nombre de balises <img> dans le contenu: {img_count}")
                # Extraire un exemple de balise d'image
                img_example = content[content.find('<img'):content.find('>', content.find('<img'))+1]
                logger.info(f"Exemple de balise image: {img_example}")
            
            # Titre du document
            title = template.get("name", "Document")
            if "title" in variables:
                title = variables["title"]
            
            # Obtenir le chemin du logo
            logo_path = None
            if self.config:
                logo_path = self.config.get("app.company_logo", "")
                
                # VÃ©rifier que le logo existe
                if logo_path and not os.path.exists(logo_path):
                    logger.warning(f"Logo spÃ©cifiÃ© mais introuvable: {logo_path}")
                    logo_path = None
            
            # GÃ©nÃ©rer le document selon le format avec gestion d'erreur
            try:
                if format_type == "pdf":
                    success = self.generate_pdf(file_path, content, title, client, company_info, logo_path)
                elif format_type == "docx":
                    success = self.generate_docx(file_path, content, title, client, company_info, logo_path)
                else:
                    success = self.generate_txt(file_path, content, title, client, company_info)
                
                if not success:
                    raise Exception(f"Ã‰chec lors de la gÃ©nÃ©ration au format {format_type}")
                
                logger.info(f"Document gÃ©nÃ©rÃ© avec succÃ¨s: {file_path}")
                return True
            
            except Exception as format_error:
                logger.error(f"Erreur lors de la gÃ©nÃ©ration au format {format_type}: {format_error}")
                
                # Solution de repli: gÃ©nÃ©rer un document texte
                try:
                    txt_path = os.path.splitext(file_path)[0] + ".txt"
                    success = self.generate_txt(txt_path, content, title, client, company_info)
                    
                    if success:
                        logger.info(f"Document texte de secours gÃ©nÃ©rÃ©: {txt_path}")
                        return True
                    else:
                        raise Exception("Ã‰chec de la gÃ©nÃ©ration du document texte de secours")
                
                except Exception as fallback_error:
                    logger.error(f"Erreur lors de la gÃ©nÃ©ration du document texte de secours: {fallback_error}")
                    
                    # DerniÃ¨re tentative: crÃ©er un document texte minimal
                    try:
                        minimal_path = os.path.splitext(file_path)[0] + ".txt"
                        with open(minimal_path, 'w', encoding='utf-8') as f:
                            f.write(f"Titre: {title}\n")
                            f.write(f"Date: {datetime.now().strftime('%Y-%m-%d')}\n")
                            f.write(f"Client: {client.get('name', '')}\n\n")
                            f.write("Une erreur est survenue lors de la gÃ©nÃ©ration du document.\n")
                            f.write("Ce document est une version de secours minimale.\n")
                        
                        logger.info(f"Document texte minimal crÃ©Ã©: {minimal_path}")
                        return True
                    
                    except Exception as minimal_error:
                        logger.error(f"Erreur lors de la crÃ©ation du document minimal: {minimal_error}")
                        return False
        
        except Exception as e:
            logger.error(f"Erreur globale lors de la gÃ©nÃ©ration du document: {e}")
            
            try:
                # Tentative ultime: crÃ©er un fichier texte d'erreur
                error_path = os.path.join(os.path.dirname(file_path), f"error_{datetime.now().strftime('%Y%m%d%H%M%S')}.txt")
                with open(error_path, 'w', encoding='utf-8') as f:
                    f.write(f"ERREUR DE GÃ‰NÃ‰RATION DE DOCUMENT\n")
                    f.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write(f"Erreur: {str(e)}\n")
                
                logger.info(f"Fichier d'erreur crÃ©Ã©: {error_path}")
                return False
            except:
                return False
    
    def generate_txt(self, file_path, content, title, client, company_info):
        """
        GÃ©nÃ¨re un document texte simple en retirant les balises HTML
        
        Args:
            file_path: Chemin du fichier Ã  crÃ©er
            content: Contenu du document (peut contenir du HTML)
            title: Titre du document
            client: Informations du client
            company_info: Informations de l'entreprise
            
        Returns:
            bool: True si rÃ©ussi, False sinon
        """
        try:
            # CrÃ©er le rÃ©pertoire parent si nÃ©cessaire
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Nettoyer le contenu HTML
            clean_content = self.strip_html(content)[0]
            
            with open(file_path, 'w', encoding='utf-8') as f:
                # En-tÃªte
                f.write(f"Titre: {title}\n\n")
                
                # Contenu
                f.write("=== CONTENU ===\n")
                f.write(clean_content)
            
            return True
        
        except Exception as e:
            logger.error(f"Erreur lors de la gÃ©nÃ©ration du fichier texte: {e}")
            return False
    
    def generate_pdf(self, file_path, content, title, client, company_info, logo_path=None):
        """
        GÃ©nÃ¨re un document PDF en retirant les balises HTML
        
        Args:
            file_path: Chemin du fichier Ã  crÃ©er
            content: Contenu du document (peut contenir du HTML)
            title: Titre du document
            client: Informations du client
            company_info: Informations de l'entreprise
            logo_path: Chemin du logo de l'entreprise
            
        Returns:
            bool: True si rÃ©ussi, False sinon
        """
        try:
            if not REPORTLAB_AVAILABLE:
                logger.error("ReportLab n'est pas installÃ©, impossible de gÃ©nÃ©rer un PDF")
                return False
            
            # CrÃ©er le rÃ©pertoire parent si nÃ©cessaire
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Nettoyer le contenu HTML et extraire les images
            clean_content, images = self.strip_html(content)
            logger.info(f"GÃ©nÃ©ration PDF: {len(images)} images extraites")
            
            # CrÃ©er un PDF
            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4
            
            # DÃ©finir la police Ã  utiliser - Helvetica par dÃ©faut
            font_name = "Helvetica"
            
            # Essayer d'utiliser une police qui supporte mieux l'Unicode si disponible
            try:
                # VÃ©rifier l'existence de DejaVuSans.ttf
                dejavu_path = None
                
                # Chemins possibles pour DejaVuSans.ttf
                possible_paths = [
                    os.path.join(os.path.dirname(__file__), 'fonts', 'DejaVuSans.ttf'),  # Dans le dossier fonts/ relatif Ã  ce script
                    os.path.join(os.path.dirname(os.path.dirname(__file__)), 'fonts', 'DejaVuSans.ttf'),  # Dans le dossier fonts/ du projet
                    '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',  # Chemin Linux commun
                    'C:\\Windows\\Fonts\\DejaVuSans.ttf',  # Chemin Windows
                    os.path.expanduser('~/Library/Fonts/DejaVuSans.ttf')  # Chemin macOS
                ]
                
                for path in possible_paths:
                    if os.path.exists(path):
                        dejavu_path = path
                        break
                
                if dejavu_path:
                    pdfmetrics.registerFont(TTFont('DejaVuSans', dejavu_path))
                    c.setFont("DejaVuSans", 10)
                    font_name = "DejaVuSans"
                    logger.info(f"Utilisation de la police DejaVuSans pour le PDF")
            except Exception as font_error:
                logger.warning(f"Impossible d'utiliser la police DejaVuSans, utilisation de Helvetica: {font_error}")
            
            # Ajouter le logo si fourni
            if logo_path and os.path.exists(logo_path):
                try:
                    # Tenter de charger le logo
                    logo_img = PILImage.open(logo_path)
                    logo_width, logo_height = logo_img.size
                    # Redimensionner si trop grand
                    max_logo_width = 150  # Taille maximale en points
                    max_logo_height = 70
                    if logo_width > max_logo_width or logo_height > max_logo_height:
                        ratio = min(max_logo_width / logo_width, max_logo_height / logo_height)
                        logo_width, logo_height = int(logo_width * ratio), int(logo_height * ratio)
                    
                    # Positionner le logo en haut Ã  droite
                    c.drawInlineImage(logo_path, width - logo_width - 50, height - 50 - logo_height, logo_width, logo_height)
                except Exception as logo_error:
                    logger.warning(f"Impossible d'ajouter le logo: {logo_error}")
            
            # DÃ©finir les marges et la largeur disponible
            margin_left = 50
            margin_right = 50
            available_width = width - margin_left - margin_right
            
            # Titre du document - avec gestion de titre long
            c.setFont(font_name + "-Bold" if font_name == "Helvetica" else font_name, 16)
            title = self._sanitize_text(title)
            
            # VÃ©rifier si le titre est trop long
            if c.stringWidth(title, font_name + "-Bold" if font_name == "Helvetica" else font_name, 16) > available_width:
                # DÃ©couper le titre en plusieurs lignes
                words = title.split(' ')
                current_line = ""
                y_position = height - 70
                
                for word in words:
                    test_line = current_line + " " + word if current_line else word
                    if c.stringWidth(test_line, font_name + "-Bold" if font_name == "Helvetica" else font_name, 16) <= available_width:
                        current_line = test_line
                    else:
                        # Dessiner la partie actuelle du titre
                        c.drawString(margin_left, y_position, current_line)
                        y_position -= 20  # Espacement pour le titre
                        current_line = word
                
                # Dessiner la derniÃ¨re partie du titre
                if current_line:
                    c.drawString(margin_left, y_position, current_line)
                    y_position -= 20
                
                # Ligne sÃ©paratrice aprÃ¨s la derniÃ¨re ligne du titre
                c.line(50, y_position - 5, width - 50, y_position - 5)
                y_position -= 20  # Espacement aprÃ¨s la ligne
            else:
                # Titre normal sur une seule ligne
                c.drawString(margin_left, height - 70, title)
                
                # Ligne sÃ©paratrice
                c.line(50, height - 85, width - 50, height - 85)
                
                y_position = height - 105
            
            # Contenu du document
            # Diviser le contenu en lignes
            lines = clean_content.split('\n')
            c.setFont(font_name, 10)
            
            font_size = 10
            image_index = 0  # Pour suivre quelle image afficher
            
            # Journaliser les lignes qui contiennent [IMAGE] pour dÃ©bogage
            image_lines = [i for i, line in enumerate(lines) if line.strip() == '[IMAGE]']
            logger.info(f"Lignes avec marqueur [IMAGE]: {image_lines}")
            
            for line in lines:
                # VÃ©rifier si c'est un marqueur d'image
                if line.strip() == '[IMAGE]' and image_index < len(images):
                    logger.info(f"Traitement de l'image {image_index+1}/{len(images)}")
                    # Tenter d'ajouter l'image
                    try:
                        img_src = images[image_index]['src']
                        logger.info(f"Source de l'image: {img_src[:50]}...")
                        # RÃ©cupÃ©rer les dimensions si prÃ©sentes
                        img_width_attr = images[image_index].get('width')
                        img_height_attr = images[image_index].get('height')
                        image_index += 1
                        
                        # VÃ©rifier si l'image est une URL ou un chemin local
                        is_url = img_src.startswith(('http://', 'https://', 'ftp://'))
                        is_temp = False  # Pour suivre si le fichier est temporaire (base64)
                        
                        # Chemin oÃ¹ tÃ©lÃ©charger/copier l'image si nÃ©cessaire
                        temp_img_path = None
                        
                        if is_url:
                            # TÃ©lÃ©charger l'image si c'est une URL
                            import tempfile
                            import urllib.request
                            
                            try:
                                # CrÃ©er un fichier temporaire
                                temp_fd, temp_img_path = tempfile.mkstemp(suffix='.jpg')
                                os.close(temp_fd)
                                
                                # TÃ©lÃ©charger l'image
                                urllib.request.urlretrieve(img_src, temp_img_path)
                                logger.info(f"Image tÃ©lÃ©chargÃ©e de {img_src} vers {temp_img_path}")
                                is_temp = True  # Marquer comme fichier temporaire
                            except Exception as url_error:
                                logger.warning(f"Erreur lors du tÃ©lÃ©chargement de l'image {img_src}: {url_error}")
                                temp_img_path = None
                        else:
                            # VÃ©rifier si c'est dÃ©jÃ  un fichier temporaire (cas des images base64)
                            if os.path.exists(img_src) and 'temp' in img_src:
                                temp_img_path = img_src
                                is_temp = True  # Marquer comme fichier temporaire
                            else:
                                # Chemin local
                                temp_img_path = img_src if os.path.isabs(img_src) else os.path.join(os.path.dirname(file_path), img_src)
                        
                        # VÃ©rifier la taille de l'image
                        img = PILImage.open(temp_img_path)
                        original_width, original_height = img.size
                        logger.info(f"PDF - Dimensions originales de l'image: {original_width}x{original_height}")
                        
                        # Par dÃ©faut, utiliser les dimensions originales
                        img_width, img_height = original_width, original_height
                        
                        # Utiliser les dimensions spÃ©cifiÃ©es si disponibles
                        if img_width_attr is not None and img_height_attr is not None:
                            logger.info(f"PDF - Dimensions spÃ©cifiÃ©es: width={img_width_attr}, height={img_height_attr}")
                            # Conversion explicite des attributs en nombres
                            try:
                                # Utiliser float puis int pour gÃ©rer les valeurs possiblement dÃ©cimales
                                img_width = int(float(img_width_attr))
                                img_height = int(float(img_height_attr))
                                
                                # S'assurer que les dimensions sont raisonnables
                                if img_width <= 0 or img_height <= 0:
                                    img_width, img_height = original_width, original_height
                                    logger.warning(f"PDF - Dimensions invalides, utilisation des dimensions d'origine: {img_width}x{img_height}")
                                else:
                                    logger.info(f"PDF - Utilisation des dimensions spÃ©cifiÃ©es aprÃ¨s conversion: {img_width}x{img_height}")
                            except Exception as dim_error:
                                logger.warning(f"PDF - Erreur lors de la conversion des dimensions: {dim_error}")
                                # Garder les dimensions d'origine en cas d'erreur
                                img_width, img_height = original_width, original_height
                        
                        # Redimensionner si trop grande
                        max_img_width = available_width  # Largeur disponible dÃ©jÃ  calculÃ©e plus haut
                        max_img_height = 300  # Points
                        
                        # S'assurer que les variables sont de type numÃ©rique pour les calculs
                        img_width = float(img_width)
                        img_height = float(img_height)
                        
                        if img_width > max_img_width or img_height > max_img_height:
                            ratio = min(max_img_width / img_width, max_img_height / img_height)
                            img_width, img_height = int(img_width * ratio), int(img_height * ratio)
                            logger.info(f"PDF - Image redimensionnÃ©e pour tenir sur la page: {img_width}x{img_height}")
                        
                        # Utiliser la position relative pour dÃ©terminer le positionnement vertical si disponible
                        relative_position = images[image_index-1].get('relative_position')
                        if relative_position is not None:
                            # Calculer la position y en fonction de la position relative dans le document
                            # Ajustement pour respecter plus fidÃ¨lement la position dans l'Ã©diteur
                            usable_height = height - 140  # Hauteur utilisable en prenant en compte les marges
                            
                            # Calcul amÃ©liorÃ©: plus la position relative est grande (vers le bas du document),
                            # plus l'image sera basse dans le PDF
                            page_height_adjustment = 0
                            
                            # Pour les images tout en haut
                            if relative_position < 0.1:
                                y_position = height - 100
                            # Pour les images tout en bas
                            elif relative_position > 0.9:
                                # VÃ©rifier s'il reste assez d'espace, sinon nouvelle page
                                if y_position - img_height < 100:
                                    c.showPage()
                                    c.setFont(font_name, font_size)
                                    y_position = height - 100
                                y_position = 150  # PrÃ¨s du bas de page
                            # Pour les autres positions
                            else:
                                # Position relative directement mappÃ©e sur la hauteur disponible
                                y_position = height - 100 - (relative_position * (height - 200))
                            
                            # S'assurer que l'image tiendra sur la page actuelle
                            if y_position - img_height < 70:
                                c.showPage()
                                c.setFont(font_name, font_size)
                                y_position = height - 100
                            
                            logger.info(f"PDF - Position relative de l'image: {relative_position}, Position y calculÃ©e: {y_position}")
                        else:
                            # Comportement par dÃ©faut si la position relative n'est pas disponible
                            # VÃ©rifier s'il reste assez d'espace sur la page
                            if y_position - img_height < 70:
                                c.showPage()
                                c.setFont(font_name, font_size)
                                y_position = height - 70
                        
                        # DÃ©terminer la position x en fonction de l'alignement
                        align = images[image_index-1].get('align', 'center')  # Par dÃ©faut, centrer l'image
                        logger.info(f"PDF - Alignement de l'image: {align}")
                        
                        if align == 'left':
                            x_pos = margin_left
                        elif align == 'right':
                            x_pos = width - margin_right - img_width
                        else:  # center par dÃ©faut
                            x_pos = margin_left + (available_width - img_width) / 2
                        
                        # Dessiner l'image avec les dimensions calculÃ©es (s'assurer qu'elles sont des nombres)
                        logger.info(f"PDF - Dessin de l'image en position ({x_pos}, {y_position - img_height}) avec dimensions {img_width}x{img_height}")
                        
                        try:
                            c.drawInlineImage(temp_img_path, x_pos, y_position - img_height, width=img_width, height=img_height)
                        except Exception as draw_error:
                            logger.error(f"PDF - Erreur lors du dessin de l'image: {draw_error}")
                            # Essayer avec les dimensions originales en cas d'erreur
                            try:
                                logger.info(f"PDF - Tentative avec dimensions originales: {original_width}x{original_height}")
                                c.drawInlineImage(temp_img_path, x_pos, y_position - original_height, 
                                                 width=original_width, height=original_height)
                            except Exception as fallback_error:
                                logger.error(f"PDF - Ã‰chec de la tentative de secours: {fallback_error}")
                        
                        # DÃ©placer le curseur y
                        y_position -= (img_height + 15)  # 15 points d'espace supplÃ©mentaire
                        
                        # Supprimer le fichier temporaire si c'en est un
                        if is_temp and temp_img_path:
                            try:
                                os.remove(temp_img_path)
                            except:
                                pass
                    except Exception as img_process_error:
                        logger.warning(f"Erreur lors du traitement de l'image: {img_process_error}")
                else:
                    # Sanitize line text
                    line = self._sanitize_text(line)
                    
                    # VÃ©rifier si la ligne est trop longue 
                    if c.stringWidth(line, font_name, font_size) > available_width:
                        # DÃ©coupage de la ligne en plusieurs lignes
                        words = line.split(' ')
                        current_line = ""
                        
                        for word in words:
                            test_line = current_line + " " + word if current_line else word
                            if c.stringWidth(test_line, font_name, font_size) <= available_width:
                                current_line = test_line
                            else:
                                # VÃ©rifier s'il reste assez d'espace sur la page
                                if y_position < 70:  # Marge de bas de page augmentÃ©e
                                    c.showPage()  # Nouvelle page
                                    c.setFont(font_name, 10)
                                    y_position = height - 70  # Marge supÃ©rieure augmentÃ©e
                                
                                # Ã‰crire la ligne actuelle
                                c.drawString(margin_left, y_position, current_line)
                                y_position -= 15  # Espacement des lignes
                                current_line = word
                        
                        # Ã‰crire la derniÃ¨re partie de la ligne
                        if current_line:
                            # VÃ©rifier s'il reste assez d'espace sur la page
                            if y_position < 70:
                                c.showPage()
                                c.setFont(font_name, 10)
                                y_position = height - 70
                            
                            c.drawString(margin_left, y_position, current_line)
                            y_position -= 15
                    else:
                        # VÃ©rifier s'il reste assez d'espace sur la page
                        if y_position < 70:  # Marge de bas de page augmentÃ©e
                            c.showPage()  # Nouvelle page
                            c.setFont(font_name, 10)
                            y_position = height - 70  # Marge supÃ©rieure augmentÃ©e
                        
                        # Dessiner la ligne
                        c.drawString(margin_left, y_position, line)
                        y_position -= 15  # Espacement des lignes
            
            # Finaliser le document
            c.showPage()
            c.save()
            
            return True
            
        except Exception as e:
            logger.error(f"Erreur lors de la gÃ©nÃ©ration du PDF: {e}")
            return False
    
    def _sanitize_text(self, text):
        """
        Nettoie le texte pour Ã©viter les problÃ¨mes avec les caractÃ¨res spÃ©ciaux dans les PDF
        
        Args:
            text: Texte Ã  nettoyer
            
        Returns:
            str: Texte nettoyÃ©
        """
        if not isinstance(text, str):
            text = str(text)
        
        # Remplacer les caractÃ¨res problÃ©matiques par leurs Ã©quivalents compatibles
        special_chars_map = {
            'â€¦': '...',
            'â€“': '-',
            'â€”': '-',
            ''': "'",
            ''': "'",
            '"': '"',
            '"': '"',
            'Â«': '"',
            'Â»': '"',
            'â€¢': '*'
        }
        
        for char, replacement in special_chars_map.items():
            text = text.replace(char, replacement)
        
        # Supprimer les caractÃ¨res non imprimables ou non supportÃ©s
        text = ''.join(c for c in text if ord(c) >= 32 or c in '\n\r\t')
        
        return text
    
    def generate_docx(self, file_path, content, title, client, company_info, logo_path=None):
        """
        GÃ©nÃ¨re un document DOCX en retirant les balises HTML
        
        Args:
            file_path: Chemin du fichier Ã  crÃ©er
            content: Contenu du document (peut contenir du HTML)
            title: Titre du document
            client: Informations du client
            company_info: Informations de l'entreprise
            logo_path: Chemin du logo de l'entreprise
            
        Returns:
            bool: True si rÃ©ussi, False sinon
        """
        try:
            if not DOCX_AVAILABLE:
                logger.error("python-docx n'est pas installÃ©, impossible de gÃ©nÃ©rer un DOCX")
                return False
            
            # CrÃ©er le rÃ©pertoire parent si nÃ©cessaire
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            
            # Nettoyer le contenu HTML et extraire les images
            clean_content, images = self.strip_html(content)
            logger.info(f"GÃ©nÃ©ration DOCX: {len(images)} images extraites")
            
            # CrÃ©er un document Word
            doc = docx.Document()
            
            # DÃ©finir la largeur maximale des marges
            for section in doc.sections:
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
                section.top_margin = Inches(1.2)
                section.bottom_margin = Inches(1.2)
            
            # Ajouter le logo si fourni
            if logo_path and os.path.exists(logo_path):
                try:
                    # Ajouter un paragraphe pour le logo
                    logo_para = doc.add_paragraph()
                    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    logo_run = logo_para.add_run()
                    logo_run.add_picture(logo_path, width=Inches(2.0))
                except Exception as logo_error:
                    logger.warning(f"Impossible d'ajouter le logo: {logo_error}")
            
            # Ajouter une police Unicode compatible Ã  l'ensemble du document
            doc_style = doc.styles['Normal']
            doc_style.font.name = 'Arial Unicode MS'
            
            # Titre du document
            title_para = doc.add_heading(title, level=1)
            
            # Ligne sÃ©paratrice
            doc.add_paragraph("_______________________________________________________________")
            
            # Contenu
            # Diviser le contenu en paragraphes
            paragraphs = clean_content.split('\n\n')
            image_index = 0  # Pour suivre quelle image afficher
            
            # Journaliser les paragraphes qui contiennent [IMAGE] pour dÃ©bogage
            image_paragraphs = [i for i, para in enumerate(paragraphs) if para.strip() == '[IMAGE]']
            logger.info(f"Paragraphes avec marqueur [IMAGE]: {image_paragraphs}")
            
            for paragraph_text in paragraphs:
                if paragraph_text.strip() == '[IMAGE]' and image_index < len(images):
                    logger.info(f"Traitement de l'image DOCX {image_index+1}/{len(images)}")
                    # Tenter d'ajouter l'image
                    try:
                        img_src = images[image_index]['src']
                        logger.info(f"Source de l'image DOCX: {img_src[:50]}...")
                        # RÃ©cupÃ©rer les dimensions si prÃ©sentes
                        img_width_attr = images[image_index].get('width')
                        img_height_attr = images[image_index].get('height')
                        image_index += 1
                        
                        # VÃ©rifier si l'image est une URL ou un chemin local
                        is_url = img_src.startswith(('http://', 'https://', 'ftp://'))
                        is_temp = False  # Pour suivre si le fichier est temporaire (base64)
                        
                        # Chemin oÃ¹ tÃ©lÃ©charger/copier l'image si nÃ©cessaire
                        temp_img_path = None
                        
                        if is_url:
                            # TÃ©lÃ©charger l'image si c'est une URL
                            import tempfile
                            import urllib.request
                            
                            try:
                                # CrÃ©er un fichier temporaire
                                temp_fd, temp_img_path = tempfile.mkstemp(suffix='.jpg')
                                os.close(temp_fd)
                                
                                # TÃ©lÃ©charger l'image
                                urllib.request.urlretrieve(img_src, temp_img_path)
                                logger.info(f"Image tÃ©lÃ©chargÃ©e de {img_src} vers {temp_img_path}")
                                is_temp = True  # Marquer comme fichier temporaire
                            except Exception as url_error:
                                logger.warning(f"Erreur lors du tÃ©lÃ©chargement de l'image {img_src}: {url_error}")
                                temp_img_path = None
                        else:
                            # VÃ©rifier si c'est dÃ©jÃ  un fichier temporaire (cas des images base64)
                            if os.path.exists(img_src) and 'temp' in img_src:
                                temp_img_path = img_src
                                is_temp = True  # Marquer comme fichier temporaire
                            else:
                                # Chemin local
                                temp_img_path = img_src if os.path.isabs(img_src) else os.path.join(os.path.dirname(file_path), img_src)
                        
                        # Ajouter l'image si disponible
                        if temp_img_path and os.path.exists(temp_img_path):
                            # VÃ©rifier la position relative de l'image pour dÃ©terminer son placement dans le document
                            relative_position = images[image_index-1].get('relative_position')
                            
                            # Si la position relative est dÃ©finie, essayer de crÃ©er un paragraphe Ã  un emplacement approximativement Ã©quivalent
                            if relative_position is not None:
                                logger.info(f"DOCX - Position relative de l'image: {relative_position}")
                                
                                # Meilleur positionnement des images basÃ© sur leur position relative dans l'Ã©diteur
                                if relative_position < 0.1:
                                    # Image tout en haut du document
                                    logger.info("Placement de l'image en haut du document")
                                    # Supprimer tous les paragraphes vides prÃ©cÃ©dents pour s'assurer que l'image est au dÃ©but
                                    # (Aucune action spÃ©ciale nÃ©cessaire)
                                elif relative_position > 0.9:
                                    # Image tout en bas
                                    logger.info("Placement de l'image tout en bas du document")
                                    # Ajouter plusieurs paragraphes vides pour pousser l'image vers le bas
                                    for _ in range(10):
                                        doc.add_paragraph("")
                                elif relative_position > 0.75:
                                    # Image dans le dernier quart
                                    logger.info("Placement de l'image dans le dernier quart du document")
                                    for _ in range(7):
                                        doc.add_paragraph("")
                                elif relative_position > 0.5:
                                    # Image dans la seconde moitiÃ©
                                    logger.info("Placement de l'image dans la seconde moitiÃ© du document")
                                    for _ in range(5):
                                        doc.add_paragraph("")
                                elif relative_position > 0.25:
                                    # Image dans le second quart
                                    logger.info("Placement de l'image dans le second quart du document")
                                    for _ in range(2):
                                        doc.add_paragraph("")
                                else:
                                    # Image dans le premier quart (mais pas tout en haut)
                                    logger.info("Placement de l'image dans le premier quart du document")
                                    doc.add_paragraph("")
                            
                            # CrÃ©er un paragraphe pour l'image
                            img_para = doc.add_paragraph()
                            
                            # DÃ©terminer l'alignement
                            align = images[image_index-1].get('align', 'center')  # Par dÃ©faut, centrer l'image
                            logger.info(f"DOCX - Alignement de l'image: {align}")
                            
                            # Appliquer l'alignement au paragraphe
                            if align == 'left':
                                img_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            elif align == 'right':
                                img_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:  # center par dÃ©faut
                                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            img_run = img_para.add_run()
                            
                            # Calculer la largeur en pouces
                            img = PILImage.open(temp_img_path)
                            original_width, original_height = img.size
                            
                            # Par dÃ©faut, largeur maximum de 5 pouces
                            default_width_inches = 5.0
                            
                            # Dimensions finales Ã  utiliser (en pouces)
                            img_width_inches = default_width_inches
                            
                            # Si les dimensions sont spÃ©cifiÃ©es, les utiliser
                            if img_width_attr is not None:
                                try:
                                    # Forcer la conversion en float puis int
                                    img_width_px = int(float(img_width_attr))
                                    
                                    # Ne pas accepter de valeurs trop petites
                                    if img_width_px < 10:
                                        img_width_px = original_width
                                        logger.warning(f"Largeur d'image trop petite, utilisation de l'original: {img_width_px}px")
                                        
                                        # Convertir pixels en pouces (Ã  96 DPI, standard pour Word)
                                        img_width_inches = img_width_px / 96.0
                                        logger.info(f"Largeur d'image spÃ©cifiÃ©e en pixels: {img_width_px} -> {img_width_inches} pouces")
                                        
                                        # Limites raisonnables pour la largeur
                                        if img_width_inches > 6.0:  # max 6 pouces
                                            img_width_inches = 6.0
                                        
                                    except Exception as e:
                                        logger.warning(f"Erreur lors de la conversion de la largeur: {e}")
                                        # En cas d'erreur, utiliser une largeur par dÃ©faut
                                        img_width_inches = min(original_width / 96.0, 6.0)
                                    
                                # Journaliser les valeurs pour dÃ©bogage
                                logger.info(f"Image: largeur originale={original_width}px, largeur spÃ©cifiÃ©e={img_width_attr}px")
                                logger.info(f"Image insÃ©rÃ©e dans DOCX avec largeur={img_width_inches} pouces")
                                
                                # Ajouter l'image avec la largeur calculÃ©e
                                try:
                                    # Utiliser Inches pour une conversion prÃ©cise
                                    img_run.add_picture(temp_img_path, width=Inches(img_width_inches))
                                    logger.info(f"Image DOCX ajoutÃ©e avec succÃ¨s: {img_width_inches} pouces")
                                except Exception as pic_error:
                                    logger.warning(f"Erreur lors de l'ajout de l'image au document: {pic_error}")
                                    # Tentative de secours avec largeur fixe
                                    try:
                                        img_run.add_picture(temp_img_path, width=Inches(4.0))
                                    except:
                                        logger.error("Impossible d'ajouter l'image mÃªme avec dimensions fixes")
                                
                                # Supprimer le fichier temporaire si c'en est un
                                if is_temp and temp_img_path:
                                    try:
                                        os.remove(temp_img_path)
                                    except:
                                        pass
                    except Exception as img_error:
                        logger.warning(f"Erreur lors de l'ajout de l'image: {img_error}")
                elif paragraph_text.strip():
                    # Pour Word, nous n'avons pas besoin de sanitizer le texte comme pour PDF
                    # car DOCX gÃ¨re bien les caractÃ¨res Unicode
                    p = doc.add_paragraph(paragraph_text)
                    # Limiter la largeur des paragraphes en ajustant les retraits
                    p.paragraph_format.first_line_indent = Inches(0)
                    p.paragraph_format.left_indent = Inches(0)
                    p.paragraph_format.right_indent = Inches(0)
                    
                    # S'assurer que la police prend en charge l'Unicode
                    for run in p.runs:
                        run.font.name = 'Arial Unicode MS'
            
            # Enregistrer le document
            doc.save(file_path)
            
            return True
            
        except Exception as e:
            logger.error(f"Erreur lors de la gÃ©nÃ©ration du DOCX: {e}")
            return False
